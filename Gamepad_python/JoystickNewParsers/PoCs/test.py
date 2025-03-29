from docx import Document
from datetime import datetime, timedelta

# Initialize the Document
doc = Document()

# Starting date for the schedule (11th September)
start_date = datetime(2023, 9, 11)
weekend_days = ['Saturday', 'Sunday']

# Define the goals for each day
goals = [
    # Week 1
    "C++ Basics Review: OOP, Classes, Inheritance, Polymorphism",
    "C++ Basics Review: Practice OOP Concepts",
    "Constructors and Destructors in C++",
    "Memory Management and Pointers in C++",
    "Introduction to STL (Standard Template Library)",
    "Sorting Algorithms: Bubble Sort, Quicksort, Mergesort",
    "Rest and Review Week 1",
    # Week 2
    "Graphs Basics: Adjacency Lists and Matrices",
    "BFS & DFS Graph Algorithms",
    "Shortest Path Algorithms: Dijkstra's and Bellman-Ford",
    "Minimum Spanning Tree (MST) Algorithms",
    "Implement Graph Algorithms: Practice",
    "Review Graph Algorithms",
    "Rest and Review Week 2",
    # Week 3
    "Introduction to Path Planning",
    "A* Algorithm and Practice",
    "A* Algorithm: Advanced Concepts and Practice",
    "Probabilistic Roadmaps (PRM): Theory and Implementation",
    "Probabilistic Roadmaps (PRM): Advanced Practice",
    "Grid-Based Search Algorithms",
    "Rest and Review Week 3",
    # Week 4
    "Rapidly-Exploring Random Trees (RRT): Concepts",
    "Rapidly-Exploring Random Trees (RRT): Practice",
    "Hybrid A* Algorithm: Concepts",
    "Hybrid A* Algorithm: Implementation",
    "Sampling-Based Algorithms: RRT* and Variants",
    "Obstacle Detection and Avoidance Techniques",
    "Rest and Review Week 4",
    # Week 5
    "Introduction to Control Systems",
    "PID Control: Theory and Applications",
    "PID Control: Implementation and Tuning",
    "Control System Design: Feedback and Feedforward Control",
    "Review and Practice Control Systems",
    "Rest and Review Week 5",
    # Week 6
    "Forward Kinematics: Concepts and Practice",
    "Inverse Kinematics: Theory and Practice",
    "Robot Dynamics: Forces and Motion",
    "Implementing Motion Control with Dynamics",
    "Rest and Review Week 6",
    # Week 7
    "Introduction to Sensors: LIDAR, GPS, Cameras",
    "Kalman Filter for State Estimation: Theory and Practice",
    "Particle Filter for Localization: Concepts",
    "Implement Particle Filter: Practice",
    "Sensor Fusion: Combining Multiple Sensors",
    "Implement a Simple Sensor Fusion System",
    "Rest and Review Week 7",
    # Week 8
    "Combining Path Planning with Control Systems",
    "Building a Complete Path Planning System",
    "Testing and Debugging the System",
    "Optimize Path Planning Algorithms",
    "Final Project Wrap-Up and Reflections",
]

# Function to generate the schedule, skipping weekends
def generate_schedule(start_date, goals):
    schedule = []
    current_date = start_date
    for week, goal in enumerate(goals, 1):
        # Skip weekends
        while current_date.strftime('%A') in weekend_days:
            current_date += timedelta(days=1)
        # Add the current goal with the corresponding date
        schedule.append((current_date.strftime('%A, %d %B %Y'), goal))
        current_date += timedelta(days=1)
    return schedule

# Generate the schedule
schedule = generate_schedule(start_date, goals)

# Write the schedule to the document
doc.add_heading('Study Plan Schedule', 0)

current_week = 1
for date, goal in schedule:
    # Add a heading for the week
    if "Rest" in goal:
        doc.add_paragraph(f"---\n**Rest Day:** {goal}")
    else:
        doc.add_paragraph(f"**{date}**\n**Goal:** {goal}\n")

# Save the document
doc.save('Study_Plan_Schedule.docx')

